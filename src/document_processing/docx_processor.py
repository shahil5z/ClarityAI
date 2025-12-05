import os
from typing import List
from docx import Document as DocxDocument  # Renamed to avoid conflict
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from src.utils.config import Config
from src.utils.logger import setup_logger

logger = setup_logger()

class DocxProcessor:
    def __init__(self):
        self.chunk_size = Config.CHUNK_SIZE
        self.chunk_overlap = Config.CHUNK_OVERLAP
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
    
    def process_docx(self, file_path: str) -> List[LangchainDocument]:
        """Process a Word document and return a list of document chunks."""
        try:
            logger.info(f"Processing Word document: {file_path}")
            
            # Extract text from Word document
            text = self._extract_text_from_docx(file_path)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents
            documents = [
                LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": os.path.basename(file_path),
                        "page": i // 10  # Approximate page number
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            logger.info(f"Created {len(documents)} document chunks from {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a Word document."""
        try:
            doc = DocxDocument(file_path)  # Using the renamed import
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {str(e)}")
            raise
